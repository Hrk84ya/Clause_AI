import re
from pathlib import Path

import chardet
import fitz  # PyMuPDF
from docx import Document as DocxDocument


def extract_text(file_path: str, mime_type: str) -> tuple[str, int, int]:
    """
    Extract text from a document file.
    Returns (full_text, page_count, word_count).
    """
    path = Path(file_path)

    if mime_type == "application/pdf":
        return _extract_pdf(path)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(path)
    elif mime_type == "text/plain":
        return _extract_txt(path)
    else:
        raise ValueError(f"Unsupported MIME type: {mime_type}")


def _extract_pdf(path: Path) -> tuple[str, int, int]:
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
    doc.close()

    full_text = "\n\n".join(pages)
    page_count = len(pages) if pages else doc.page_count
    word_count = len(full_text.split())

    return full_text, page_count, word_count


def _extract_docx(path: Path) -> tuple[str, int, int]:
    """Extract text from DOCX using python-docx, preserving heading structure."""
    doc = DocxDocument(str(path))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Preserve heading structure
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            paragraphs.append(f"\n{text}\n")
        else:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n".join(paragraphs)
    page_count = 1  # DOCX doesn't have a reliable page count without rendering
    word_count = len(full_text.split())

    return full_text, page_count, word_count


def _extract_txt(path: Path) -> tuple[str, int, int]:
    """Extract text from TXT with encoding detection."""
    raw_bytes = path.read_bytes()

    # Detect encoding
    detected = chardet.detect(raw_bytes)
    encoding = detected.get("encoding", "utf-8") or "utf-8"

    try:
        full_text = raw_bytes.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        full_text = raw_bytes.decode("utf-8", errors="replace")

    page_count = 1
    word_count = len(full_text.split())

    return full_text, page_count, word_count


def extract_headings(text: str) -> list[str]:
    """
    Extract section headings from text.
    Detects: numbered headings (1., 1.1, etc.), ALL CAPS lines.
    """
    headings = []
    lines = text.split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Numbered headings: "1.", "1.1", "1.1.1", "Section 1", etc.
        if re.match(r"^(\d+[\.\d]*)\s+[A-Z]", stripped):
            headings.append(stripped)
        # ALL CAPS lines (at least 6 chars, mostly uppercase)
        elif len(stripped) >= 6 and stripped.isupper():
            headings.append(stripped)
        # "ARTICLE" or "SECTION" prefixed
        elif re.match(r"^(ARTICLE|SECTION|CLAUSE)\s+", stripped, re.IGNORECASE):
            headings.append(stripped)

    return headings
